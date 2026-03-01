#!/usr/bin/env python
"""
Word document generator for Anglerphish reports.
"""

import datetime
import os
import sys
import json
import requests
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

_SURROGATE_RE = re.compile(r'[\ud800-\udfff]')

# Python-docx for creating Word (.docx) files
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.stderr.write("Error: python-docx library is required. Install with: pip install python-docx\n")
    sys.exit(1)

# Optional user-agents parsing
try:
    from user_agents import parse
    USER_AGENTS_AVAILABLE = True
except ImportError:
    USER_AGENTS_AVAILABLE = False
    sys.stderr.write("Warning: user_agents library not available. Install with: pip install user-agents\n")

# Optional GeoIP for IP location
try:
    import geoip2.database
    GEOIP_AVAILABLE = True
    GEOIP_DB_PATHS = [
        'GeoLite2-City.mmdb',
        os.path.join(os.path.dirname(__file__), 'GeoLite2-City.mmdb'),
        os.path.join(os.path.expanduser('~'), 'GeoLite2-City.mmdb'),
        '/usr/share/GeoIP/GeoLite2-City.mmdb',
        '/usr/local/share/GeoIP/GeoLite2-City.mmdb'
    ]
    GEOIP_DB_PATH = next((p for p in GEOIP_DB_PATHS if os.path.isfile(p)), None)
    GEOIP_READER = geoip2.database.Reader(GEOIP_DB_PATH) if GEOIP_DB_PATH else None
except ImportError:
    GEOIP_AVAILABLE = False
    GEOIP_READER = None

def set_cell_background(cell, fill_hex):
    """Set the cell shading (background color). fill_hex is a string like '4472C4' (no '#')."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_table_borders_color(table, border_hex='FFFFFF', border_size='4', border_style='single'):
    """
    Turn all table borders (outer + inner) to the given hex color.
      • border_hex: 'RRGGBB' (no ‘#’)
      • border_size: width in eighths of a point (string)
      • border_style: e.g. 'single', 'double', etc.
    """
    tbl = table._tbl

    # 1) Grab or create <w:tblPr>
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # 2) Remove any existing <w:tblBorders>
    for old in list(tblPr.findall(qn('w:tblBorders'))):
        tblPr.remove(old)

    # 3) Build new borders block
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'),   border_style)
        elem.set(qn('w:sz'),    border_size)
        elem.set(qn('w:color'), border_hex)
        borders.append(elem)

    # 4) Append in place
    tblPr.append(borders)

def style_table(
    table,
    header_fill='4472C4',
    row_fill='D9D9D9',
    alt_row_fill='BFBFBF',
    border_hex='FFFFFF',
    header_font_hex='FFFFFF'
):
    # 1) Set borders
    set_table_borders_color(table, border_hex=border_hex)

    # 2) Header styling
    if header_fill is not None and header_font_hex is not None:
            r, g, b = (int(header_font_hex[i:i+2], 16) for i in (0, 2, 4))
            for cell in table.rows[0].cells:
                set_cell_background(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(r, g, b)

    # 3) Alternating row styling
    for i, row in enumerate(table.rows[1:], start=1):
        fill = row_fill if i % 2 else alt_row_fill
        for cell in row.cells:
            set_cell_background(cell, fill)


def clean_text(s):
    """Strip invalid UTF-16 surrogates from a string for Word compatibility."""
    if isinstance(s, str):
        return _SURROGATE_RE.sub('', s)
    return str(s) if s is not None else ''


# Helper: format date
def format_date(date_string):
    """Date formatter for mixed ISO/UTC/custom formats."""
    if not date_string or date_string == 'N/A':
        return 'N/A'

    try:
        # Strip Z or timezone if present
        if 'Z' in date_string:
            date_string = date_string.rstrip('Z')

        # Handle microseconds with too many digits
        if 'T' in date_string and '.' in date_string:
            date_part, time_part = date_string.split('T')
            if '+' in time_part or '-' in time_part:
                # Has timezone info, split carefully
                if '+' in time_part:
                    time_main, tz = time_part.split('+')
                    tz = '+' + tz
                else:
                    time_main, tz = time_part.rsplit('-', 1)
                    tz = '-' + tz
            else:
                time_main = time_part
                tz = ''

            if '.' in time_main:
                time_clean, micro = time_main.split('.')
                micro = micro[:6]  # truncate microseconds
                time_main = f"{time_clean}.{micro}"

            date_string = f"{date_part}T{time_main}{tz}"

        # Convert to datetime
        dt = datetime.datetime.fromisoformat(date_string)
        return dt.strftime('%Y-%m-%d %H:%M:%S')

    except Exception:
        return date_string  # fallback for invalid or malformed strings

# Helper: determine if campaign/result is SMS or Generic
def is_sms_campaign(campaign):
    """Check if a campaign is SMS-based"""
    return campaign.get('type') == 'sms'

def is_generic_campaign(campaign):
    """Check if a campaign is Generic (landing page only, no email/SMS)"""
    return campaign.get('type') == 'generic'

def is_generic_result(result):
    """Check if a result is for a generic campaign link (no email/phone)"""
    # Generic results have first_name as link name but no email/phone
    return not result.get('email') and not result.get('phone')

def is_sms_result(result):
    """Check if a result is for an SMS target"""
    # Check if it's explicitly marked as SMS
    if result.get('sms_target'):
        return True
    # Or if it has a phone but no email
    if result.get('phone') and not result.get('email'):
        return True
    return False

def get_contact_field(result):
    """Get the contact identifier (email or phone) for a result"""
    if is_sms_result(result):
        return result.get('phone', 'Unknown')
    return result.get('email', 'Unknown')

def anonymize_phone(phone):
    """Anonymize a phone number by masking the middle digits"""
    if not phone or phone == 'Unknown':
        return phone
    # Remove all non-digit characters for processing
    digits = ''.join(c for c in phone if c.isdigit())
    if len(digits) < 4:
        return phone  # Too short to anonymize meaningfully
    # Show first 2 and last 2 digits, mask the rest
    return phone[:2] + '*' * (len(digits) - 4) + phone[-2:]

def anonymize_email(email):
    """Anonymize an email address"""
    if not email or email == 'Unknown' or '@' not in email:
        return email
    local, domain = email.split('@', 1)
    if len(local) <= 2:
        return email
    return local[0] + '*' * (len(local) - 2) + local[-1] + '@' + domain

# Helper: extract UA
def extract_user_agent(result):
    b = result.get('browser') if isinstance(result.get('browser'), dict) else None
    if b:
        return b.get('user-agent')
    d = result.get('details') if isinstance(result.get('details'), dict) else None
    if d and isinstance(d.get('browser'), dict):
        return d['browser'].get('user-agent')
    return None

def has_user_interaction(result):
    """Check if user had any interaction with the phishing content"""
    return any([
        result.get('opened'),
        result.get('clicked'),
        result.get('submitted_data'),
        result.get('reported'),
        result.get('replied'),
        result.get('status') in ('Email Opened', 'Clicked Link', 'Submitted Data', 'Email Replied', 'Email Reported')
    ])

def extract_browser_info(result, check_interaction=True):
    """Extract browser info. Returns N/A if no user interaction and check_interaction is True."""
    if check_interaction and not has_user_interaction(result):
        return "N/A"
    ua = extract_user_agent(result)
    if ua and USER_AGENTS_AVAILABLE:
        try:
            parsed = parse(ua)
            return f"{parsed.browser.family} {parsed.browser.version_string}"
        except:
            return ua
    return "Unknown"

def extract_os_info(result, check_interaction=True):
    """Extract OS info. Returns N/A if no user interaction and check_interaction is True."""
    if check_interaction and not has_user_interaction(result):
        return "N/A"
    ua = extract_user_agent(result)
    if ua and USER_AGENTS_AVAILABLE:
        try:
            parsed = parse(ua)
            return f"{parsed.os.family} {parsed.os.version_string}"
        except:
            return ua
    return "Unknown"

def extract_ip_address(result, check_interaction=True):
    """Extract IP address. Returns N/A if no user interaction and check_interaction is True."""
    if check_interaction and not has_user_interaction(result):
        return "N/A"
    ip = result.get('ip')
    if ip:
        return ip
    d = result.get('details') if isinstance(result.get('details'), dict) else {}
    b = d.get('browser') if isinstance(d.get('browser'), dict) else None
    return b.get('address') if b and b.get('address') else 'Unknown'

def is_private_ip(ip):
    if not ip: return False
    if ip.startswith(('10.','192.168.','127.','169.254.','::1','fc00:','fd')):
        return True
    if ip.startswith('172.'):
        try:
            sec = int(ip.split('.')[1])
            return 16 <= sec <=31
        except: pass
    return False

def get_ip_location(ip, result=None):
    """
    Try GeoLite2 → FreeIPAPI → Gophish lat/lon → Unknown.
    """
    if not ip or ip == 'Unknown':
        return 'Unknown'

    # 1) Private ranges
    if is_private_ip(ip):
        return 'Internal Network'

    # 2) GeoLite2 database
    if GEOIP_AVAILABLE and GEOIP_READER:
        try:
            resp = GEOIP_READER.city(ip)
            city = resp.city.name or ""
            country = resp.country.name or ""
            if city or country:
                return f"{city}, {country}"
        except Exception:
            pass

    # 3) FreeIPAPI fallback
    if requests:
        try:
            r = requests.get(f"https://freeipapi.com/api/json/{ip}", timeout=5)
            if r.ok:
                data = r.json()
                # FreeIPAPI fields: cityName, regionName, countryName :contentReference[oaicite:1]{index=1}
                city   = data.get('cityName')   or data.get('city')   or ""
                region = data.get('regionName') or data.get('region') or ""
                country= data.get('countryName') or data.get('country') or ""
                loc = ", ".join(p for p in (city, region, country) if p)
                if loc:
                    return loc
        except Exception:
            pass

    # 4) Gophish coordinates fallback
    if result and isinstance(result, dict):
        lat = result.get('latitude')
        lon = result.get('longitude')
        if lat is not None and lon is not None:
            return f"{lat}, {lon}"

    return 'Unknown'

# helper: given a list of result-dicts, returns a dict of named delta lists
def _parse_iso(s: str) -> datetime.datetime:
    """Strip a trailing 'Z', truncate any fractional‐seconds to 6 digits,
       and preserve a '+HH:MM' or '-HH:MM' offset for fromisoformat."""
    s = s.rstrip('Z')
    # split off the timezone offset if present
    tz = ''
    m = re.search(r'([+-]\d{2}:\d{2})$', s)
    if m:
        tz = m.group(1)
        s  = s[: -len(tz)]
    # truncate fractions to 6 digits
    if '.' in s:
        date_part, frac = s.split('.', 1)
        # keep only digits, then pad/truncate to 6
        frac = ''.join(ch for ch in frac if ch.isdigit())[:6].ljust(6, '0')
        s = f"{date_part}.{frac}"
    return datetime.datetime.fromisoformat(s + tz)

# Helper function to extract payload data
def extract_payload_from_dict(data_dict):
    result = ""
    if isinstance(data_dict, dict) and 'payload' in data_dict and isinstance(data_dict['payload'], dict):
        for key, value in data_dict['payload'].items():
            if isinstance(value, list) and value:
                result += f"{key}: \"{value[0]}\"\n"
            else:
                result += f"{key}: \"{value}\"\n"
    return result

def collect_all_deltas(results):
    click_deltas = []
    subm_deltas  = []
    reply_deltas = []
    report_deltas= []

    mapping = [
        ('clicked_time',    click_deltas),
        ('submitted_time',  subm_deltas),
        ('replied_time',    reply_deltas),
        ('reported_time',   report_deltas),
    ]

    for r in results:
        dt_str = r.get('send_date')
        if not dt_str:
            continue
        try:
            delivered = _parse_iso(dt_str)
        except Exception:
            continue

        for field, lst in mapping:
            ts = r.get(field)
            if not ts:
                continue
            try:
                t = _parse_iso(ts)
                lst.append((t - delivered).total_seconds() / 60)
            except Exception:
                pass

    # failure = minimum of click/submit/reply per user
    failure_deltas = []
    for r in results:
        times = []
        for fld in ('clicked_time','submitted_time','replied_time'):
            ts = r.get(fld)
            if not ts:
                continue
            try:
                t = _parse_iso(ts)
                d = (t - _parse_iso(r['send_date'])).total_seconds() / 60
                times.append(d)
            except:
                pass
        if times:
            failure_deltas.append(min(times))

    return {
        'Click':        click_deltas,
        'Submit Data':  subm_deltas,
        'Reply':        reply_deltas,
        'Report':       report_deltas,
        'Failure':      failure_deltas,
    }

def generate_word_document(data, output_path, include_toc=True, gdpr_options=None):
    """
    Generate a Word (.docx) report from GoPhish campaign data.
    Note: include_toc parameter is kept for backward compatibility but not used.
    """
    # Validate data structure
    if not isinstance(data, dict) or 'campaigns' not in data or not data['campaigns']:
        sys.stderr.write("Error: Invalid or empty campaign data\n")
        return False
    
    # Preprocess: inject browser, IP, and action flags from timeline events into each result
    for campaign in data['campaigns']:
        # Build results map using both email and phone as keys
        results_map = {}
        for r in campaign.get('results', []):
            email_key = r.get('email')
            phone_key = r.get('phone')
            if email_key:
                results_map[email_key] = r
            if phone_key:
                results_map[phone_key] = r
        
        for event in campaign.get('timeline', []):
            contact = event.get('email')  # Field name is 'email' but may contain phone for SMS
            ts = event.get('time')
            if contact not in results_map:
                continue

            result = results_map[contact]
            msg = event.get('message')

            # Record the moment the message was sent/delivered (email or SMS)
            if msg in ('Email Sent', 'Email Delivered', 'SMS Sent', 'SMS Delivered'):
                result['send_date'] = ts

            # Record every action flag
            if msg == 'Email Opened':
                result['opened'] = True
                result['opened_time'] = ts
            elif msg == 'Clicked Link':
                result['clicked'] = True
                result['clicked_time'] = ts
            elif msg == 'Submitted Data':
                result['submitted_data'] = True
                result['submitted_time'] = ts
            elif msg == 'Email Replied':
                result['replied'] = True
                result['replied_time'] = ts
            elif msg == 'Email Reported':
                result['reported'] = True
                result['reported_time'] = ts

            # Inject browser & IP once
            details = event.get('details')
            if isinstance(details, str):
                try:
                    details = json.loads(details)
                except json.JSONDecodeError:
                    details = {}
            
            result['details'] = details
            if 'payload' in details:
                result['payload'] = details['payload']

            browser = details.get('browser') if isinstance(details, dict) else None
            if isinstance(browser, dict):
                result['browser'] = browser
                if 'address' in browser:
                    result['ip'] = browser['address']

    # Create document
    doc = Document()
    from docx.enum.style  import WD_STYLE_TYPE
    from docx.shared      import Pt, RGBColor, Inches

    # — Define a custom bullet style —  
    bullet_style = doc.styles.add_style('CustomBullet', WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.base_style = doc.styles['List Bullet']

    # Font tweaks
    f = bullet_style.font
    f.name = 'Calibri'
    f.size = Pt(11)
    f.color.rgb = RGBColor(0x2F, 0x4F, 0x4F)    # dark slate gray

    # Paragraph tweaks
    pf = bullet_style.paragraph_format
    pf.left_indent        = Inches(0.3)
    pf.first_line_indent  = Inches(-0.15)
    pf.space_before       = Pt(2)
    pf.space_after        = Pt(2)
    title = doc.add_heading("Anglerphish Campaign Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para = doc.add_paragraph(f"Generated on: {datetime.datetime.now():%Y-%m-%d %H:%M:%S}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add privacy notice if anonymization is enabled
    if gdpr_options and (gdpr_options.get('anonymize_emails', False) or gdpr_options.get('anonymize_ips', False)):
        doc.add_heading("Privacy Notice", level=2)
        privacy_para = doc.add_paragraph(
            "This report has been generated with personal data obfuscated in accordance with privacy requirements. "
            "Email addresses/phone numbers and IP addresses have been partially masked to protect individual privacy."
        )
        privacy_para.style = 'Intense Quote'
        doc.add_paragraph()  # Add spacing after the privacy notice

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This report provides an analysis of the phishing campaign(s) conducted using the Anglerphish platform, "
        "including performance metrics and key user interactions."
    )

    # Flatten results
    all_results = []
    all_results_non_generic = []  # Results from email/SMS campaigns only (for user-based stats)
    in_progress = False
    has_email_campaigns = False
    has_sms_campaigns = False
    has_generic_campaigns = False
    
    # Separate tracking for generic campaign events (count every event, not unique users)
    generic_total_links = 0
    generic_total_clicks = 0
    generic_total_submissions = 0
    
    for c in data['campaigns']:
        results = c.get('results', [])
        all_results.extend(results)
        
        if c.get('status') not in ('Complete', 'Completed'):
            in_progress = True
        
        # Track campaign types
        if is_generic_campaign(c):
            has_generic_campaigns = True
            generic_total_links += len(results)
            # For generic campaigns, count EVERY event from timeline (not unique)
            timeline = c.get('timeline', [])
            for event in timeline:
                if event.get('message') == 'Clicked Link':
                    generic_total_clicks += 1
                elif event.get('message') == 'Submitted Data':
                    generic_total_submissions += 1
        elif is_sms_campaign(c):
            has_sms_campaigns = True
            all_results_non_generic.extend(results)
        else:
            has_email_campaigns = True
            all_results_non_generic.extend(results)

    # Build a map of unique users with their highest action across ALL campaigns
    # Action levels: 4=submitted, 3=clicked, 2=opened, 1=sent
    user_highest_action = {}  # email -> highest action level
    user_campaigns_failed = {}  # email -> set of campaign names where user failed
    
    for c in data['campaigns']:
        campaign_name = c.get('name', 'Unnamed')
        for r in c.get('results', []):
            email = r.get('email') or r.get('phone', 'Unknown')
            
            # Determine action level for this result
            if r.get('submitted_data') or r.get('status') == 'Submitted Data':
                action_level = 4
            elif r.get('clicked') or r.get('status') == 'Clicked Link':
                action_level = 3
            elif r.get('opened') or r.get('status') == 'Email Opened':
                action_level = 2
            else:
                action_level = 1  # Email Sent only
            
            # Track highest action per user
            if email not in user_highest_action:
                user_highest_action[email] = action_level
            else:
                user_highest_action[email] = max(user_highest_action[email], action_level)
            
            # Track campaigns where user failed (clicked, submitted, or replied)
            if r.get('clicked') or r.get('replied') or r.get('submitted_data'):
                if email not in user_campaigns_failed:
                    user_campaigns_failed[email] = set()
                user_campaigns_failed[email].add(campaign_name)
    
    # Count unique users per action level (mutually exclusive)
    total_unique_users = len(user_highest_action)
    submitted = sum(1 for level in user_highest_action.values() if level == 4)
    clicked = sum(1 for level in user_highest_action.values() if level == 3)
    opened = sum(1 for level in user_highest_action.values() if level == 2)
    sent_only = sum(1 for level in user_highest_action.values() if level == 1)
    
    # For replied and reported, count unique users who did these actions (across all campaigns)
    users_replied = set()
    users_reported = set()
    for r in all_results:
        email = r.get('email') or r.get('phone', 'Unknown')
        if r.get('replied') or r.get('status') == 'Email Replied':
            users_replied.add(email)
        if r.get('reported') or r.get('status') == 'Email Reported':
            users_reported.add(email)
    replied = len(users_replied)
    reported = len(users_reported)
    
    # Use total_unique_users for percentage calculations
    total = total_unique_users

    def pct(p): return round((p/total)*100) if total else 0

    doc.add_heading("Overall Campaign Results", level=2)
    
    # Add explanatory note about deduplication
    doc.add_paragraph(
        "The following table shows unique users across all campaigns. Each user is counted once, "
        "categorized by their highest-risk action (e.g., if a user clicked in Campaign A but submitted "
        "data in Campaign B, they are counted under 'Data Submitted')."
    )
    
    # Build row_data dynamically based on campaign types
    row_data = [
        ("Data Submitted", submitted, pct(submitted), 'F05B4F'),  # Red
        ("Clicks", clicked, pct(clicked), 'F39C12'),  # Orange
    ]
    
    # Only include Email-specific metrics if there are email campaigns
    if has_email_campaigns:
        row_data.append(("Email Opened", opened, pct(opened), 'F9BF3B'))  # Amber/Yellow
        row_data.append(("Email Replied", replied, pct(replied), 'E67E22'))  # Dark orange
        row_data.append(("Reported", reported, pct(reported), '45D6EF'))  # Light blue
    
    row_data.append(("Total", total, '', 'D9D9D9'))  # Gray - no percentage for total
    
    # Create table with appropriate number of rows
    stats_tbl = doc.add_table(rows=len(row_data) + 1, cols=3, style='Table Grid')
    
    # Header row
    hdr = stats_tbl.rows[0].cells
    hdr[0].text = 'Category'
    hdr[1].text = 'Users'
    hdr[2].text = 'Percentage'
    
    # Style header row with blue background and white text
    for cell in hdr:
        set_cell_background(cell, '4472C4')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
    
    # Add data and style the cells (using the dynamically created row_data from above)
    for i, (label, count, percentage, color) in enumerate(row_data, start=1):
        row = stats_tbl.rows[i].cells
        
        # Category column
        row[0].text = label
        set_cell_background(row[0], color)
        for p in row[0].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
        
        # Users column
        row[1].text = str(count)
        set_cell_background(row[1], 'E7E6E6')  # Light gray
        
        # Percentage column
        row[2].text = f"{percentage}%" if percentage != '' else ''
        set_cell_background(row[2], 'E7E6E6')  # Light gray
    
    # Set table borders to white
    set_table_borders_color(stats_tbl, border_hex='FFFFFF')

    # Add Generic Campaign Results section if there are generic campaigns
    if has_generic_campaigns:
        doc.add_paragraph()  # spacing
        doc.add_heading("Generic Campaign Results (Event Counts)", level=2)
        doc.add_paragraph(
            "Generic campaigns count every event (not unique users) since links can be shared "
            "and accessed by multiple anonymous users. The following shows total engagement events."
        )
        
        generic_row_data = [
            ("Total Links", generic_total_links, '', '1ABC9C'),  # Green
            ("Total Clicks", generic_total_clicks, '', 'F39C12'),  # Orange
            ("Total Submissions", generic_total_submissions, '', 'F05B4F'),  # Red
        ]
        
        generic_tbl = doc.add_table(rows=len(generic_row_data) + 1, cols=2, style='Table Grid')
        
        # Header row
        hdr = generic_tbl.rows[0].cells
        hdr[0].text = 'Metric'
        hdr[1].text = 'Count'
        
        for cell in hdr:
            set_cell_background(cell, '4472C4')
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True
        
        for i, (label, count, _, color) in enumerate(generic_row_data, start=1):
            row = generic_tbl.rows[i].cells
            row[0].text = label
            set_cell_background(row[0], color)
            for p in row[0].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True
            
            row[1].text = str(count)
            set_cell_background(row[1], 'E7E6E6')
        
        set_table_borders_color(generic_tbl, border_hex='FFFFFF')

    if in_progress:
        doc.add_paragraph(
            "Note: Includes in-progress campaigns; data may change as they complete.",
            style='Intense Quote'
        )

    # Add per-scenario results table
    doc.add_paragraph()  # spacing
    scenario_tbl = doc.add_table(rows=1, cols=6, style='Table Grid')
    
    # Header row with specific colors
    hdr = scenario_tbl.rows[0].cells
    hdr[0].text = 'Scenario'
    hdr[1].text = 'Targets\nFailed'
    hdr[2].text = 'Reported\nPhishing'
    hdr[3].text = 'Total Targets'
    hdr[4].text = 'Fail %\nPercentage'
    hdr[5].text = 'Report %\nPercentage'
    
    # Header colors: Scenario (blue), Targets Failed (red), Reported Phishing (cyan), rest (blue)
    header_colors = ['4472C4', 'C55A11', '00B0F0', '4472C4', '4472C4', '4472C4']
    
    for idx, (cell, color) in enumerate(zip(hdr, header_colors)):
        set_cell_background(cell, color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
    
    # Add data rows for each campaign/scenario
    for campaign in data['campaigns']:
        name = campaign.get('name', 'N/A')
        results = campaign.get('results', [])
        total_targets = len(results)
        is_generic_camp = is_generic_campaign(campaign)
        
        # For Generic campaigns, count events from timeline (not unique users)
        if is_generic_camp:
            # Count total click events from timeline
            targets_failed = sum(1 for e in campaign.get('timeline', []) if e.get('message') == 'Clicked Link')
            reported_count = 0  # Generic campaigns don't have reporting
            # For Generic, show click rate based on number of links
            fail_pct = round((targets_failed / total_targets) * 100, 2) if total_targets else 0
            report_pct = 0
        else:
            # Count failures (clicked, submitted data, or replied) - unique users
            # Use email or phone for SMS campaigns
            targets_failed = len({
                r.get('email') or r.get('phone', 'Unknown')
                for r in results
                if r.get('clicked') or r.get('replied') or r.get('submitted_data')
            })
            
            # Count reported
            reported_count = sum(1 for r in results if r.get('reported') or r.get('status')=='Email Reported')
            
            # Calculate percentages
            fail_pct = round((targets_failed / total_targets) * 100, 2) if total_targets else 0
            report_pct = round((reported_count / total_targets) * 100, 2) if total_targets else 0
        
        # Add row
        row = scenario_tbl.add_row().cells
        row[0].text = clean_text(name)
        # For Generic campaigns, show "Total Clicks" label context
        row[1].text = str(targets_failed) + (" clicks" if is_generic_camp else "")
        row[2].text = str(reported_count) if not is_generic_camp else "N/A"
        row[3].text = str(total_targets) + (" links" if is_generic_camp else "")
        row[4].text = f"{fail_pct}%"
        row[5].text = f"{report_pct}%" if not is_generic_camp else "N/A"
        
        # Alternate row coloring - light gray and white
        row_num = len(scenario_tbl.rows) - 1  # Current row number (excluding header)
        fill_color = 'E7E6E6' if row_num % 2 else 'FFFFFF'
        for cell in row:
            set_cell_background(cell, fill_color)
    
    # Set table borders
    set_table_borders_color(scenario_tbl, border_hex='FFFFFF')

    # Add "Repeat Offenders" table - users who failed in more than one campaign
    repeat_offenders = {email: campaigns for email, campaigns in user_campaigns_failed.items() if len(campaigns) > 1}
    
    if repeat_offenders and len(data['campaigns']) > 1:
        doc.add_paragraph()  # spacing
        doc.add_heading("Repeat Offenders", level=2)
        doc.add_paragraph(
            "The following users failed (clicked, submitted data, or replied) in more than one campaign. "
            "These users may require additional security awareness training."
        )
        
        # Create table with Email/Phone, # of Campaigns Failed, Campaign Names
        repeat_tbl = doc.add_table(rows=1, cols=3, style='Table Grid')
        
        # Header row
        hdr = repeat_tbl.rows[0].cells
        hdr[0].text = 'Email/Phone'
        hdr[1].text = 'Campaigns Failed'
        hdr[2].text = 'Campaign Names'
        
        # Style header row
        for cell in hdr:
            set_cell_background(cell, 'C55A11')  # Orange/Red for warning
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True
        
        # Add data rows sorted by number of campaigns failed (descending)
        for email, campaigns in sorted(repeat_offenders.items(), key=lambda x: -len(x[1])):
            row = repeat_tbl.add_row().cells
            
            # Apply anonymization if enabled
            display_email = email
            if gdpr_options and gdpr_options.get('anonymize_emails'):
                if '@' in email:
                    display_email = anonymize_email(email)
                else:
                    display_email = anonymize_phone(email)
            
            row[0].text = clean_text(display_email)
            row[1].text = str(len(campaigns))
            row[2].text = clean_text(', '.join(sorted(campaigns)))
            
            # Alternate row coloring
            row_num = len(repeat_tbl.rows) - 1
            fill_color = 'E7E6E6' if row_num % 2 else 'FFFFFF'
            for cell in row:
                set_cell_background(cell, fill_color)
        
        # Set table borders
        set_table_borders_color(repeat_tbl, border_hex='FFFFFF')

    doc.add_heading("Key Takeaways", level=1)
    # — Phish-Prone Percentage (unique failures) —
    unique_failures = len({
        r.get('email')
        for r in all_results
        if r.get('clicked') or r.get('replied') or r.get('submitted_data')
    })
    ppp = pct(unique_failures)
    # use the List Bullet style instead of Intense Quote
    para = doc.add_paragraph(
        f"Phish-Prone Percentage: {unique_failures} unique users out of {total} "
        f"({ppp}%) failed by engaging in risky behavior (clicked, replied, or submitted data).",
        style='List Bullet'
    )
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT


    # — Execution Summary: speed metrics bullets —
    actions = collect_all_deltas(all_results)
    bullets = []
    def pct(lst, thresh):
        return round(100 * sum(1 for d in lst if d <= thresh) / len(lst)) if lst else 0

    # near-instant risk
    inst = pct(actions['Failure'], 3)
    if inst:
        bullets.append(f"{inst}% of failures occurred in under 3 min (near-instant risk)")

    # reporting stats - fix the logic to properly count actual reports
    # Count total users who actually reported (from the reported flag, regardless of timing)
    total_reported = sum(1 for r in all_results if r.get('reported') or r.get('status')=='Email Reported')
    
    # Calculate never reported percentage based on actual reporting behavior
    never_reported_pct = round(((total - total_reported) / total) * 100) if total else 0
    bullets.append(f"{never_reported_pct}% of users never reported phishing")

    # For timing-based metrics, only use reports with valid timing data
    reps = actions['Report']  # This contains timing deltas for reports with timing data
    if reps:
        # Only show timing metrics when we have timing data
        never24 = round(100 * sum(1 for d in reps if d > 1440) / len(reps))
        bullets.append(f"{never24}% of users who reported did so after 24 h")
        
        median = sorted(reps)[len(reps)//2]
        bullets.append(f"Median time-to-report: {round(median)} min")
        fast = pct(reps, 3)
        if fast:
            bullets.append(f"{fast}% of users reported phishing within 3 min")
    elif total_reported > 0:
        # We have reports but no timing data
        bullets.append(f"Timing data not available for {total_reported} reports")

    # now emit each bullet with the List Bullet style
    for text in bullets:
        para = doc.add_paragraph(text, style='List Bullet')
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()  # spacing


     # — Results Summary per Campaign —
    doc.add_heading("Campaign Results", level=1)

    # helper for per‐campaign percentages
    def pct_local(part, whole):
        return round((part/whole)*100) if whole else 0

    for c in data['campaigns']:
        results = c.get('results', [])
        sent      = len(results)
        
        # Calculate statistics using mutually exclusive counts (highest action only)
        # Hierarchy: Data Submitted > Clicked Link > Email Opened > Email Sent
        subm_c = sum(1 for r in results if r.get('submitted_data') or r.get('status')=='Submitted Data')
        clicked_c = sum(1 for r in results if (r.get('clicked') or r.get('status')=='Clicked Link') and not (r.get('submitted_data') or r.get('status')=='Submitted Data'))
        opened_c = sum(1 for r in results if (r.get('opened') or r.get('status')=='Email Opened') and not (r.get('clicked') or r.get('status')=='Clicked Link'))
        # Email Sent = users who never opened (ensures total = submitted + clicked + opened + sent)
        sent_only_c = sent - subm_c - clicked_c - opened_c
        
        # Replied and reported are separate actions, not part of the main funnel
        replied_c = sum(1 for r in results if r.get('replied') or r.get('status')=='Email Replied')
        rep_c     = sum(1 for r in results if r.get('reported') or r.get('status')=='Email Reported')

        # — Campaign heading —  
        doc.add_heading(f"Campaign: {c.get('name', 'Unnamed')}", level=2)
        raw_status = c.get('status', 'N/A')
        
        # If the status is not 'Complete' or 'Completed', show progress as a percentage of sent emails 
        if raw_status not in ('Complete', 'Completed'):
            sent = len(c.get('results', []))
            total_recip = c.get('total_recipients', sent)
            if total_recip:
                pct = round(sent / total_recip * 100)
                status_display = f"{raw_status} ({pct}% – {sent}/{total_recip} sent)"
            else:
                status_display = raw_status
        else:
            status_display = raw_status

        # — Campaign Info Table (metadata only) —
        doc.add_heading("Campaign Information", level=3)
        
        # Build info fields based on campaign type
        is_sms = is_sms_campaign(c)
        is_generic = is_generic_campaign(c)
        
        # Determine campaign type display
        if is_generic:
            campaign_type_display = "Generic (Landing Page Only)"
        elif is_sms:
            campaign_type_display = "SMS"
        else:
            campaign_type_display = "Email"
        
        info_fields = [
            ("Campaign ID", str(c.get('id', 'N/A'))),
            ("Campaign Type", campaign_type_display),
            ("Created Date", format_date(c.get('created_date', 'N/A'))),
            ("Launch Date", format_date(c.get('launch_date', 'N/A'))),
            ("Completion Date", format_date(c.get('completed_date', 'N/A'))),
            ("Status", status_display),
        ]
        
        # Add type-specific fields
        if is_sms:
            # SMS campaign fields
            sms_template = c.get('sms_template', {})
            info_fields.extend([
                ("Sender ID", c.get('sms', {}).get('from', 'N/A')),
                ("SMS Provider", c.get('sms', {}).get('provider', 'N/A')),
            ])
        elif not is_generic:
            # Email campaign fields (skip for Generic campaigns)
            info_fields.extend([
                ("Email Subject", c.get('template_details', {}).get('subject', 'N/A')),
                ("Envelope Sender", c.get('template_details', {}).get('envelope_sender', 'N/A')),
            ])
        
        # Common fields for both types
        info_fields.extend([
            ("Phishing URL", c.get('phish_url', 'N/A')),
            ("URL Parameter", c.get('urlparam', 'rid')),
            ("Redirect URL", c.get('page_details', {}).get('redirect_url', 'N/A')),
            ("Data Captured", "Yes" if c.get('page_details', {}).get('capture_credentials', 0) else "No"),
            ("Passwords Stored", "Yes" if c.get('page_details', {}).get('capture_passwords', 0) else "No"),
        ])

        info_table = doc.add_table(rows=len(info_fields), cols=2, style='Table Grid')
        
        for i, (label, val) in enumerate(info_fields):
            info_table.rows[i].cells[0].text = clean_text(label)
            info_table.rows[i].cells[1].text = clean_text(val)
            
            # Color the left column cell with blue
            left_cell = info_table.rows[i].cells[0]
            set_cell_background(left_cell, '4472C4')
            
            # Make the text white and bold
            for p in left_cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True
            
            # Set the right column to grey
            right_cell = info_table.rows[i].cells[1]
            set_cell_background(right_cell, 'D9D9D9')
            
        # Set table borders to white
        set_table_borders_color(info_table, border_hex='FFFFFF')

        doc.add_paragraph()  # spacing

        # — Campaign Results Table (matching Overall Campaign Results format) —
        # For Generic campaigns, use event-based counts from timeline instead of unique users
        if is_generic:
            doc.add_heading("Campaign Event Summary", level=3)
            doc.add_paragraph(
                "Generic campaigns count every event (not unique users) since links can be shared "
                "and accessed by multiple anonymous visitors."
            )
            
            # Count events from timeline for this specific campaign
            generic_clicks = sum(1 for e in c.get('timeline', []) if e.get('message') == 'Clicked Link')
            generic_submissions = sum(1 for e in c.get('timeline', []) if e.get('message') == 'Submitted Data')
            
            # Create simple 2-column table for Generic campaigns (Metric | Count)
            generic_event_tbl = doc.add_table(rows=4, cols=2, style='Table Grid')
            
            # Header row
            hdr = generic_event_tbl.rows[0].cells
            hdr[0].text = 'Metric'
            hdr[1].text = 'Count'
            
            for cell in hdr:
                set_cell_background(cell, '4472C4')
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
            
            # Data rows
            generic_metrics = [
                ("Total Links", sent, '1ABC9C'),
                ("Total Clicks", generic_clicks, 'F39C12'),
                ("Total Submissions", generic_submissions, 'F05B4F'),
            ]
            
            for i, (label, count, color) in enumerate(generic_metrics, start=1):
                row = generic_event_tbl.rows[i].cells
                row[0].text = label
                set_cell_background(row[0], color)
                for p in row[0].paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
                
                row[1].text = str(count)
                set_cell_background(row[1], 'E7E6E6')
            
            set_table_borders_color(generic_event_tbl, border_hex='FFFFFF')
            doc.add_paragraph()  # spacing
            
            # Skip the standard campaign_row_data table for generic
            campaign_row_data = None
        else:
            doc.add_heading("Campaign Unique Results Summary", level=3)
            
            # Build row_data dynamically based on campaign type
            campaign_row_data = [
                ("Data Submitted", subm_c, pct_local(subm_c, sent), 'F05B4F'),  # Red
                ("Clicks", clicked_c, pct_local(clicked_c, sent), 'F39C12'),  # Orange
            ]
            
            # Only include Email-specific metrics for email campaigns
            if not is_sms:
                campaign_row_data.append(("Email Opened", opened_c, pct_local(opened_c, sent), 'F9BF3B'))  # Amber/Yellow
                campaign_row_data.append(("Email Sent", sent_only_c, pct_local(sent_only_c, sent), '1ABC9C'))  # Green
                campaign_row_data.append(("Email Replied", replied_c, pct_local(replied_c, sent), 'E67E22'))  # Dark orange
                campaign_row_data.append(("Reported", rep_c, pct_local(rep_c, sent), '45D6EF'))  # Light blue
            else:
                # For SMS campaigns, add "SMS Sent" row
                campaign_row_data.append(("SMS Sent", sent_only_c, pct_local(sent_only_c, sent), '1ABC9C'))  # Green
            
            campaign_row_data.append(("Total", sent, '', 'D9D9D9'))  # Gray - no percentage for total
        
        # Only create results table for non-Generic campaigns
        if campaign_row_data is not None:
            # Create table with appropriate number of rows
            results_tbl = doc.add_table(rows=len(campaign_row_data) + 1, cols=3, style='Table Grid')
            
            # Header row
            hdr = results_tbl.rows[0].cells
            hdr[0].text = 'Category'
            hdr[1].text = 'Users'
            hdr[2].text = 'Percentage'
            
            # Style header row with blue background and white text
            for cell in hdr:
                set_cell_background(cell, '4472C4')
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
            
            # Add data and style the cells (using the dynamically created campaign_row_data)
            for i, (label, count, percentage, color) in enumerate(campaign_row_data, start=1):
                row = results_tbl.rows[i].cells
                
                # Category column
                row[0].text = label
                set_cell_background(row[0], color)
                for p in row[0].paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
                
                # Users column
                row[1].text = str(count)
                set_cell_background(row[1], 'E7E6E6')  # Light gray
                
                # Percentage column
                row[2].text = f"{percentage}%" if percentage != '' else ''
                set_cell_background(row[2], 'E7E6E6')  # Light gray
            
            # Set table borders to white
            set_table_borders_color(results_tbl, border_hex='FFFFFF')

            doc.add_paragraph()  # spacing

        # — Campaign Results Summary Table (Cumulative/Hierarchical counts) —
        # Skip this table for Generic campaigns (already shown event counts above)
        if not is_generic:
            doc.add_heading("Campaign Results Summary", level=3)
            
            # Calculate cumulative counts (each action includes users who did further actions)
            # These are NOT mutually exclusive - they're cumulative/hierarchical
            # In cumulative mode: opened includes clicked/submitted, clicked includes submitted
            total_submitted = sum(1 for r in results if r.get('submitted_data') or r.get('status')=='Submitted Data')
            # Clicked = clicked OR submitted (submitting implies clicking)
            total_clicked = sum(1 for r in results if r.get('clicked') or r.get('status')=='Clicked Link' or r.get('submitted_data') or r.get('status')=='Submitted Data')
            # Opened = opened OR clicked OR submitted (clicking implies opening for cumulative)
            total_opened = sum(1 for r in results if r.get('opened') or r.get('status')=='Email Opened' or r.get('clicked') or r.get('status')=='Clicked Link' or r.get('submitted_data') or r.get('status')=='Submitted Data')
            total_replied = sum(1 for r in results if r.get('replied') or r.get('status')=='Email Replied')
            total_reported = sum(1 for r in results if r.get('reported') or r.get('status')=='Email Reported')
            
            # Build cumulative row_data dynamically based on campaign type
            cumulative_row_data = [
                ("Data Submitted", total_submitted, pct_local(total_submitted, sent), 'F05B4F'),  # Red
                ("Clicked Link", total_clicked, pct_local(total_clicked, sent), 'F39C12'),  # Orange
            ]
            
            # Only include Email-specific metrics for email campaigns
            if not is_sms:
                cumulative_row_data.append(("Email Opened", total_opened, pct_local(total_opened, sent), 'F9BF3B'))  # Amber/Yellow
                cumulative_row_data.append(("Email Sent", sent, pct_local(sent, sent), '1ABC9C'))  # Green (100%)
                cumulative_row_data.append(("Email Replied", total_replied, pct_local(total_replied, sent), 'E67E22'))  # Dark orange
                cumulative_row_data.append(("Reported", total_reported, pct_local(total_reported, sent), '45D6EF'))  # Light blue
            else:
                # For SMS campaigns
                cumulative_row_data.append(("SMS Sent", sent, pct_local(sent, sent), '1ABC9C'))  # Green (100%)
            
            # Create table with appropriate number of rows
            cumulative_tbl = doc.add_table(rows=len(cumulative_row_data) + 1, cols=3, style='Table Grid')
            
            # Header row
            hdr = cumulative_tbl.rows[0].cells
            hdr[0].text = 'Category'
            hdr[1].text = 'Users'
            hdr[2].text = 'Percentage'
            
            # Style header row with blue background and white text
            for cell in hdr:
                set_cell_background(cell, '4472C4')
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
            
            # Add data and style the cells
            for i, (label, count, percentage, color) in enumerate(cumulative_row_data, start=1):
                row = cumulative_tbl.rows[i].cells
                
                # Category column
                row[0].text = label
                set_cell_background(row[0], color)
                for p in row[0].paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
                
                # Users column
                row[1].text = str(count)
                set_cell_background(row[1], 'E7E6E6')  # Light gray
                
                # Percentage column
                row[2].text = f"{percentage}%" if percentage != '' else ''
                set_cell_background(row[2], 'E7E6E6')  # Light gray
            
            # Set table borders to white
            set_table_borders_color(cumulative_tbl, border_hex='FFFFFF')

            doc.add_paragraph()  # spacing

        # — User Interaction Summary Table —
        if results:
            # Dynamic heading based on campaign type
            if is_generic:
                doc.add_heading("Link Summary", level=3)
            else:
                doc.add_heading("User Interaction Summary", level=3)
            
            # Dynamic table structure based on campaign type
            if is_generic:
                # Generic campaigns: Link Name, Total Clicks, Total Submissions (counts per link from timeline)
                num_cols = 3
                tbl = doc.add_table(rows=1, cols=num_cols, style='Table Grid')
                hdr = tbl.rows[0].cells
                hdr[0].text = 'Link Name'
                hdr[1].text = 'Total Clicks'
                hdr[2].text = 'Total Submissions'
                
                header_colors = [
                    '4472C4',  # Link Name
                    'F39C12',  # Total Clicks
                    'F05B4F',  # Total Submissions
                ]
                
                # Build per-link counts from timeline
                # For Generic campaigns, timeline events use a tracking identifier in the 'email' field
                # We need to match this back to results by trying multiple strategies
                link_click_counts = {}
                link_submission_counts = {}
                link_click_events = {}  # Store actual events for Per Link Details
                link_submission_events = {}
                timeline = c.get('timeline', [])
                
                # Build multiple lookup maps to help match timeline events to results
                result_by_email = {}
                result_by_id = {}
                for r in results:
                    # Map by email if present
                    email = r.get('email', '')
                    if email:
                        result_by_email[email] = r
                    # Also map by ID (as string) for fallback matching
                    rid = str(r.get('id', ''))
                    if rid:
                        result_by_id[rid] = r
                
                # For Generic campaigns, we also need to track by result position/order
                # since timeline events may reference results differently
                result_by_index = {i: r for i, r in enumerate(results)}
                
                for event in timeline:
                    link_id = event.get('email', '')  # For Generic, 'email' field contains link identifier
                    msg = event.get('message')
                    
                    # Find the matching result using multiple strategies
                    matching_result = None
                    if link_id in result_by_email:
                        matching_result = result_by_email[link_id]
                    elif link_id in result_by_id:
                        matching_result = result_by_id[link_id]
                    
                    if msg == 'Clicked Link':
                        link_click_counts[link_id] = link_click_counts.get(link_id, 0) + 1
                        if link_id not in link_click_events:
                            link_click_events[link_id] = []
                        link_click_events[link_id].append(event)
                        # Also set clicked flag on the result if we can find it
                        if matching_result:
                            matching_result['clicked'] = True
                            matching_result['clicked_time'] = event.get('time')
                            # Inject browser/IP info from event
                            details = event.get('details')
                            if isinstance(details, str):
                                try:
                                    details = json.loads(details)
                                except json.JSONDecodeError:
                                    details = {}
                            if details:
                                matching_result['details'] = details
                                browser = details.get('browser') if isinstance(details, dict) else None
                                if isinstance(browser, dict):
                                    matching_result['browser'] = browser
                                    if 'address' in browser:
                                        matching_result['ip'] = browser['address']
                    elif msg == 'Submitted Data':
                        link_submission_counts[link_id] = link_submission_counts.get(link_id, 0) + 1
                        if link_id not in link_submission_events:
                            link_submission_events[link_id] = []
                        link_submission_events[link_id].append(event)
                        # Also set submitted_data flag on the result if we can find it
                        if matching_result:
                            matching_result['submitted_data'] = True
                            matching_result['submitted_time'] = event.get('time')
                            # Inject details/payload from event
                            details = event.get('details')
                            if isinstance(details, str):
                                try:
                                    details = json.loads(details)
                                except json.JSONDecodeError:
                                    details = {}
                            if details:
                                matching_result['details'] = details
                                if 'payload' in details:
                                    matching_result['payload'] = details['payload']
                
                for idx, cell in enumerate(hdr):
                    fill = header_colors[idx]
                    set_cell_background(cell, fill)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.bold = True
                
                # Populate rows for Generic campaigns (per-link counts)
                for r in results:
                    row_cells = tbl.add_row().cells
                    link_name = r.get('first_name', 'Unknown Link')
                    # For Generic campaigns, timeline events use the RId as identifier,
                    # which is serialized as "id" in the JSON (not "email")
                    link_id = r.get('id', '')  # Use 'id' (RId) to match with timeline events
                    
                    row_cells[0].text = clean_text(link_name)
                    row_cells[1].text = str(link_click_counts.get(link_id, 0))
                    row_cells[2].text = str(link_submission_counts.get(link_id, 0))
                    
                    # Store event references in the result for Per Link Details
                    r['_click_events'] = link_click_events.get(link_id, [])
                    r['_submission_events'] = link_submission_events.get(link_id, [])
                
                style_table(tbl,
                        header_fill=None,
                        row_fill='D9D9D9',
                        alt_row_fill='BFBFBF',
                        border_hex='FFFFFF',
                        header_font_hex='FFFFFF')
                
                doc.add_paragraph()  # spacing
                
            elif is_sms:
                # SMS campaigns: Phone Number, Clicked, Data Submitted, OS, Browser
                num_cols = 5
                tbl = doc.add_table(rows=1, cols=num_cols, style='Table Grid')
                hdr = tbl.rows[0].cells
                hdr[0].text = 'Phone Number'
                hdr[1].text = 'Clicked'
                hdr[2].text = 'Data Submitted'
                hdr[3].text = 'OS'
                hdr[4].text = 'Browser'
                
                header_colors = [
                    '4472C4',  # Phone number
                    'F39C12',  # Clicked
                    'F05B4F',  # Data Submitted
                    '4472C4',  # OS
                    '4472C4',  # Browser
                ]
            else:
                # Email campaigns: Email Address, Opened, Clicked, Data Submitted, Replied, Reported, OS, Browser
                num_cols = 8
                tbl = doc.add_table(rows=1, cols=num_cols, style='Table Grid')
                hdr = tbl.rows[0].cells
                hdr[0].text = 'Email Address'
                hdr[1].text = 'Opened'
                hdr[2].text = 'Clicked'
                hdr[3].text = 'Data Submitted'
                hdr[4].text = 'Replied'
                hdr[5].text = 'Reported'
                hdr[6].text = 'OS'
                hdr[7].text = 'Browser'
                
                header_colors = [
                    '4472C4',  # Email address
                    'F9BF3B',  # Opened
                    'F39C12',  # Clicked
                    'F05B4F',  # Data Submitted
                    'E67E22',  # Replied
                    '45D6EF',  # Reported
                    '4472C4',  # OS
                    '4472C4',  # Browser
                ]
            
            # Style header row (only for non-generic campaigns, generic already handled above)
            if not is_generic:
                for idx, cell in enumerate(hdr):
                    fill = header_colors[idx]
                    set_cell_background(cell, fill)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.bold = True

            # populate rows (only for non-generic campaigns, generic already handled above)
            if not is_generic:
                for r in results:
                    row_cells = tbl.add_row().cells
                    
                    # Get contact field dynamically based on result type
                    contact = get_contact_field(r)
                    
                    # Apply GDPR anonymization if enabled
                    if gdpr_options and gdpr_options.get('anonymize_emails'):
                        if is_sms_result(r):
                            # Anonymize phone number for SMS results
                            contact = anonymize_phone(contact)
                        else:
                            # Anonymize email for email results
                            contact = anonymize_email(contact)
                    
                    row_cells[0].text = clean_text(contact)
                
                    if is_sms:
                        # SMS campaigns: Clicked, Data Submitted, OS, Browser
                        cell = row_cells[1]; cell.text = ""
                        run = cell.paragraphs[0].add_run('✓' if r.get('clicked') else '✘')
                        if r.get('clicked'): run.bold = True

                        cell = row_cells[2]; cell.text = ""
                        run = cell.paragraphs[0].add_run('✓' if r.get('submitted_data') else '✘')
                        if r.get('submitted_data'): run.bold = True

                        row_cells[3].text = clean_text(extract_os_info(r))
                        row_cells[4].text = clean_text(extract_browser_info(r))
                    else:
                        # Email campaigns: Opened, Clicked, Data Submitted, Replied, Reported, OS, Browser
                        cell = row_cells[1]; cell.text = ""
                        run = cell.paragraphs[0].add_run('✓' if r.get('opened') else '✘')
                        if r.get('opened'): run.bold = True

                        cell = row_cells[2]; cell.text = ""
                        run = cell.paragraphs[0].add_run('✓' if r.get('clicked') else '✘')
                        if r.get('clicked'): run.bold = True

                        cell = row_cells[3]; cell.text = ""
                        run = cell.paragraphs[0].add_run('✓' if r.get('submitted_data') else '✘')
                        if r.get('submitted_data'): run.bold = True

                        cell = row_cells[4]; cell.text = ""
                        run = cell.paragraphs[0].add_run('✓' if r.get('replied') else '✘')
                        if r.get('replied'): run.bold = True

                        cell = row_cells[5]; cell.text = ""
                        run = cell.paragraphs[0].add_run('✓' if r.get('reported') else '✘')
                        if r.get('reported'): run.bold = True

                        row_cells[6].text = clean_text(extract_os_info(r))
                        row_cells[7].text = clean_text(extract_browser_info(r))
            
            style_table(tbl,
                    header_fill=None,
                    row_fill='D9D9D9',
                    alt_row_fill='BFBFBF',
                    border_hex='FFFFFF',
                    header_font_hex='FFFFFF')

            doc.add_paragraph()  # add a blank line for spacing

        # Dynamic heading based on campaign type
        if is_generic:
            doc.add_heading("Per Link Details", level=3)
        else:
            doc.add_heading("Per User Details", level=3)
        
        for r in results:
            # Only include users/links with any action
            if not any([r.get('opened'), r.get('clicked'), r.get('submitted_data'), r.get('reported'), r.get('replied')]):
                continue

            # Get contact/link field dynamically and apply anonymization
            if is_generic:
                contact = r.get('first_name', 'Unknown Link')
                contact_label = "Link Name"
                sent_label = "Link Created"
            else:
                contact = get_contact_field(r)
                contact_label = "Phone Number" if is_sms_result(r) else "Email"
                sent_label = "SMS Sent" if is_sms_result(r) else "Email Sent"
            
            if gdpr_options and gdpr_options.get('anonymize_emails'):
                if is_sms_result(r):
                    # Anonymize phone number for SMS results
                    contact = anonymize_phone(contact)
                else:
                    # Anonymize email for email results
                    contact = anonymize_email(contact)
            
            # Header: display real name if available, else contact
            name = r.get('full_name') or contact
            doc.add_heading(clean_text(name), level=4)
            
            # User Info Summary (condensed) - dynamic based on campaign type
            send_date = format_date(r['send_date']) if r.get('send_date') else 'N/A'
            
            # Check if user had any interaction - if no interaction (email only sent, never opened/clicked), IP should be N/A
            has_interaction = any([r.get('opened'), r.get('clicked'), r.get('submitted_data'), r.get('reported'), r.get('replied')])
            
            if has_interaction:
                ip = extract_ip_address(r)
                # Apply IP anonymization if requested
                if gdpr_options and gdpr_options.get('anonymize_ips') and ip != 'Unknown':
                    # Simple IP anonymization: mask last octet for IPv4
                    if '.' in ip and ip.count('.') == 3:
                        ip = '.'.join(ip.split('.')[:-1]) + '.***'
                location = get_ip_location(ip, result=r)
            else:
                ip = 'N/A'
                location = 'N/A'
            browser = extract_browser_info(r)
            os = extract_os_info(r)
            
            # Create user summary info table
            summary_tbl = doc.add_table(rows=6, cols=2, style='Table Grid')
            summary_data = [
                (contact_label, clean_text(contact)),
                (sent_label, send_date),
                ("IP Address", ip),
                ("Location", location),
                ("Browser", browser),
                ("Operating System", os),
            ]
            
            for i, (label, value) in enumerate(summary_data):
                summary_tbl.rows[i].cells[0].text = label
                summary_tbl.rows[i].cells[1].text = clean_text(value)
                
                # Style left column
                left_cell = summary_tbl.rows[i].cells[0]
                set_cell_background(left_cell, '4472C4')
                for p in left_cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
                
                # Style right column
                right_cell = summary_tbl.rows[i].cells[1]
                set_cell_background(right_cell, 'E7E6E6')
            
            set_table_borders_color(summary_tbl, border_hex='FFFFFF')
            doc.add_paragraph()  # spacing
            
            # User Actions Timeline Table
            doc.add_heading("User Activity Timeline" if not is_generic else "Link Activity Timeline", level=5)
            
            # Collect all actions with timestamps
            actions_list = []
            
            # For Generic campaigns, show ALL events from the stored event lists
            if is_generic:
                # Get all click events for this link
                for event in r.get('_click_events', []):
                    event_time = event.get('time', '')
                    # Extract browser/IP info from event details
                    details = event.get('details')
                    if isinstance(details, str):
                        try:
                            details = json.loads(details)
                        except json.JSONDecodeError:
                            details = {}
                    
                    browser_info = ""
                    ip_info = ""
                    if details and isinstance(details, dict):
                        browser = details.get('browser', {})
                        if browser:
                            ip_info = browser.get('address', '')
                            ua = browser.get('user-agent', '')
                            if ua and USER_AGENTS_AVAILABLE:
                                try:
                                    parsed = parse(ua)
                                    browser_info = f"{parsed.browser.family} on {parsed.os.family}"
                                except:
                                    browser_info = ua[:50] if len(ua) > 50 else ua
                    
                    detail_text = f"IP: {ip_info}" if ip_info else "Link clicked"
                    if browser_info:
                        detail_text += f" | {browser_info}"
                    
                    actions_list.append(("Link Clicked", format_date(event_time), detail_text, 'F39C12'))
                
                # Get all submission events for this link
                for event in r.get('_submission_events', []):
                    event_time = event.get('time', '')
                    details = event.get('details')
                    if isinstance(details, str):
                        try:
                            details = json.loads(details)
                        except json.JSONDecodeError:
                            details = {}
                    
                    # Extract payload
                    payload_text = ""
                    if details:
                        payload_text = extract_payload_from_dict(details)
                        # Also get IP
                        browser = details.get('browser', {}) if isinstance(details, dict) else {}
                        ip_info = browser.get('address', '') if browser else ''
                        if ip_info and not payload_text:
                            payload_text = f"IP: {ip_info}"
                        elif ip_info:
                            payload_text = f"IP: {ip_info}\n{payload_text}"
                    
                    detail = payload_text if payload_text else "Data was submitted"
                    actions_list.append(("Data Submitted", format_date(event_time), detail, 'F05B4F'))
                
                # Sort by timestamp
                actions_list.sort(key=lambda x: x[1])
            else:
                # Non-generic campaigns: show single timestamps per action type
                if r.get('opened_time'):
                    actions_list.append(("Email Opened", format_date(r['opened_time']), "User viewed the email", 'F9BF3B'))
                
                if r.get('clicked_time'):
                    actions_list.append(("Link Clicked", format_date(r['clicked_time']), "User clicked the phishing link", 'F39C12'))
                
                if r.get('submitted_time'):
                    # Extract payload for display
                    payload_text = ""
                    if 'details' in r and isinstance(r['details'], str):
                        try:
                            details = json.loads(r['details'])
                            payload_text = extract_payload_from_dict(details)
                        except json.JSONDecodeError:
                            pass
                    
                    if not payload_text and 'details' in r and isinstance(r['details'], dict):
                        payload_text = extract_payload_from_dict(r['details'])
                    
                    if not payload_text and 'browser' in r and isinstance(r['browser'], dict):
                        payload_text = extract_payload_from_dict(r['browser'])
                    
                    if not payload_text and 'payload' in r and isinstance(r['payload'], dict):
                        for key, value in r['payload'].items():
                            if isinstance(value, list) and value:
                                payload_text += f"{key}: \"{value[0]}\"\n"
                            else:
                                payload_text += f"{key}: \"{value}\"\n"
                    
                    detail = payload_text if payload_text else "Data was submitted"
                    actions_list.append(("Data Submitted", format_date(r['submitted_time']), detail, 'F05B4F'))
                
                if r.get('replied_time'):
                    actions_list.append(("Email Replied", format_date(r['replied_time']), "User replied to the phishing email", 'E67E22'))
                
                if r.get('reported_time'):
                    actions_list.append(("Phish Reported", format_date(r['reported_time']), "User reported the phishing attempt", '45D6EF'))
            
            # Create timeline table
            if actions_list:
                timeline_tbl = doc.add_table(rows=len(actions_list) + 1, cols=3, style='Table Grid')
                
                # Header row
                hdr = timeline_tbl.rows[0].cells
                hdr[0].text = 'Action'
                hdr[1].text = 'Timestamp'
                hdr[2].text = 'Details'
                
                for cell in hdr:
                    set_cell_background(cell, '4472C4')
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.bold = True
                
                # Data rows
                for i, (action, timestamp, detail, color) in enumerate(actions_list, start=1):
                    row = timeline_tbl.rows[i].cells
                    
                    # Action column with color coding
                    row[0].text = action
                    set_cell_background(row[0], color)
                    for p in row[0].paragraphs:
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.bold = True
                    
                    # Timestamp column
                    row[1].text = timestamp
                    set_cell_background(row[1], 'E7E6E6')
                    
                    # Details column
                    row[2].text = clean_text(detail)
                    set_cell_background(row[2], 'E7E6E6')
                
                set_table_borders_color(timeline_tbl, border_hex='FFFFFF')
            
            doc.add_paragraph()  # spacing between users

        # Add page break after each campaign (except the last one)
        if c != data['campaigns'][-1]:
            doc.add_page_break()

    # Add page break before Other Statistics section
    doc.add_page_break()
    doc.add_heading("Other Statistics", level=1)
    # — Time-to-Action Distribution —
    bins = [
        (float('-inf'), 3, "< 3 min"),  # Use -inf to capture d=0 and any negative deltas
        (3,    15,    "4–15 min"),
        (15,   60,    "16–60 min"),
        (60,  240,    "1–4 h"),
        (240,1440,    "4–24 h"),
        (1440, float('inf'), "> 24 h"),
    ]

    # collect all five series (Click, Submit Data, Reply, Report, Failure)
    actions = collect_all_deltas(all_results)

    doc.add_heading("Time-to-Action Distribution", level=2)
    tbl = doc.add_table(rows=1, cols=len(bins)+1, style='Table Grid')
    hdr = tbl.rows[0].cells
    hdr[0].text = "Action"
    for idx,(_,_,label) in enumerate(bins, start=1):
        hdr[idx].text = label
    
    # Style header row with blue background and white bold text
    for cell in hdr:
        set_cell_background(cell, '4472C4')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True

    def pct_bin(lst, low, high):
        return round(100 * sum(1 for d in lst if low < d <= high) / len(lst)) if lst else 0

    # Define action colors matching dashboard.js
    action_colors = {
        'Click': 'F39C12',        # Orange (for Clicked Link)
        'Submit Data': 'F05B4F',  # Red (for Submitted Data - from dashboard)
        'Reply': 'E67E22',        # Dark orange (for Email Reply)
        'Report': '45D6EF',       # Light blue (for Report)
    }

    # If you don't want "Failure" in the table, filter it out here:
    table_actions = {k:v for k,v in actions.items() if k != 'Failure'}

    for action_name, deltas in table_actions.items():
        row = tbl.add_row().cells
        
        # Set action name with color coding
        row[0].text = action_name
        action_color = action_colors.get(action_name, '4472C4')  # Default to blue if not found
        set_cell_background(row[0], action_color)
        for p in row[0].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
        
        # Set percentage values with light gray background
        for idx,(low, high, _) in enumerate(bins, start=1):
            row[idx].text = f"{pct_bin(deltas, low, high)}%"
            set_cell_background(row[idx], 'E7E6E6')  # Light gray
    
    # Set table borders to white
    set_table_borders_color(tbl, border_hex='FFFFFF')

    doc.add_paragraph()  # spacing before next section

    # Browser Stats - only for users who interacted
    doc.add_heading("Browser Statistics", level=2)
    browser_counts = {}
    for r in all_results:
        if not has_user_interaction(r):
            continue  # Skip users who never interacted
        b = extract_browser_info(r, check_interaction=False)  # Already checked
        browser_counts[b] = browser_counts.get(b,0) + 1
    tbl = doc.add_table(rows=1, cols=2, style='Table Grid')
    header_cells = tbl.rows[0].cells
    header_cells[0].text, header_cells[1].text = 'Browser','Count'
    
    # Bold the headers
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                
    for b,c in sorted(browser_counts.items(), key=lambda x:-x[1]):
        row=tbl.add_row().cells; row[0].text, row[1].text = clean_text(b),str(c)
    
    style_table(tbl)

    # OS Stats - only for users who interacted
    doc.add_heading("Operating System Statistics", level=2)
    os_counts = {}
    for r in all_results:
        if not has_user_interaction(r):
            continue  # Skip users who never interacted
        o = extract_os_info(r, check_interaction=False)  # Already checked
        os_counts[o] = os_counts.get(o,0) + 1
    tbl = doc.add_table(rows=1, cols=2, style='Table Grid')
    header_cells = tbl.rows[0].cells
    header_cells[0].text, header_cells[1].text = 'OS','Count'
    
    # Bold the headers
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                
    for o,c in sorted(os_counts.items(), key=lambda x:-x[1]):
        row=tbl.add_row().cells; row[0].text, row[1].text = clean_text(o),str(c)

    style_table(tbl)

    # IP Stats - only for users who interacted
    doc.add_heading("IP Address and Location Statistics", level=2)
    ip_counts, ip_locs = {}, {}
    for r in all_results:
        if not has_user_interaction(r):
            continue  # Skip users who never interacted
        ip = extract_ip_address(r, check_interaction=False)  # Already checked
        ip_counts[ip] = ip_counts.get(ip,0) + 1
        if ip not in ip_locs:
            ip_locs[ip] = get_ip_location(ip, result=r)
    tbl = doc.add_table(rows=1, cols=3, style='Table Grid')
    header_cells = tbl.rows[0].cells
    header_cells[0].text = 'IP Address'
    header_cells[1].text = 'Location'
    header_cells[2].text = 'Count'

    # Bold the headers as you did before
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for ip,c in sorted(ip_counts.items(), key=lambda x:-x[1]):
        row=tbl.add_row().cells; row[0].text, row[1].text, row[2].text = clean_text(ip), clean_text(ip_locs[ip]), str(c)

    style_table(tbl)

    # # Recommendations
    # doc.add_heading("Recommendations", level=1)

    # recommendations = [
    #     "Schedule targeted awareness training for users in the near-instant risk group (< 3 min failures) and those who never reported within 24 h.",
    #     "Enable advanced URL-scanning/link-rewriting on inbound email and quarantine suspicious messages.",
    #     "Enforce and monitor SPF, DKIM, and DMARC with a 'reject' policy for rogue senders.",
    #     "Provide a prominent 'Report Phish' button in the mail client and publicize its use.",
    #     "Integrate phishing-test performance into quarterly security KPIs by department.",
    #     "Plan a follow-up, more sophisticated test in 4–6 weeks for high-risk cohorts.",
    #     "Automate a weekly dashboard of PPP and time-to-action metrics; alert if PPP > 20%.",
    # ]

    # for rec in recommendations:
    #     para = doc.add_paragraph(rec, style='List Bullet')
    #     para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Save
    try:
        doc.save(output_path)
        sys.stderr.write(f"Report saved to {output_path}\n")
        return True
    except Exception as e:
        sys.stderr.write(f"Error saving document: {e}\n")
        return False

# End of script
